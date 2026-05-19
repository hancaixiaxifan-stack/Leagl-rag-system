from __future__ import annotations

import os
from dataclasses import dataclass

from docx import Document

# docx段落
@dataclass(frozen=True)
class DocxParagraph:
    para_idx: int  # 1-based index in python-docx paragraph list
    text: str

# 解析docx文件
@dataclass(frozen=True)
class ParsedDocx:
    path: str
    title: str
    paragraphs: list[DocxParagraph]  # non-empty only

# 从docx文件中猜标题
def guess_title_from_docx(doc: Document, fallback: str) -> str:
    for para in doc.paragraphs[:10]:
        t = (para.text or "").strip()
        if t and len(t) <= 60:
            return t
    return fallback

#解析docx文件
def parse_docx(path: str) -> ParsedDocx:
    doc = Document(path)
    paras: list[DocxParagraph] = []
    for i, para in enumerate(doc.paragraphs, start=1):
        t = (para.text or "").strip()
        if not t:
            continue
        # normalize common whitespace
        t = " ".join(t.split())
        paras.append(DocxParagraph(para_idx=i, text=t))

    title = guess_title_from_docx(doc, fallback=os.path.splitext(os.path.basename(path))[0])
    return ParsedDocx(path=path, title=title, paragraphs=paras)

